import os
from docx import Document


def write_code_to_word(doc, code_files):
    """
    Writes the contents of files to the Word document.

    Args:
        doc (Document): The Word document object.
        code_files (list): List of file paths to write.
    """
    doc.add_heading('Project Code:', level=1)
    for file_path in code_files:
        try:
            print(f'Processing file: {file_path}')
            doc.add_heading(f'File: {file_path}', level=2)
            with open(file_path, 'r', encoding='utf-8') as file:
                code_content = file.read()
                doc.add_paragraph(code_content)
        except Exception as e:
            doc.add_paragraph(f'Error reading {file_path}: {e}')


def write_directory_structure(doc, start_path):
    """
    Writes the directory structure to the Word document, only including files under 'hermod/hermod'.

    Args:
        doc (Document): The Word document object.
        start_path (str): Path of the directory to traverse.

    Returns:
        list: List of Python files to include their code in the Word document.
    """
    doc.add_heading(f'Project Structure ({start_path}):', level=1)
    print(f'Traversing directory: {start_path}')
    code_files = []
    for root, dirs, files in os.walk(start_path):
        if 'Hermod/src' in root.replace('\\', '/'):
            level = root.replace(start_path, '').count(os.sep)
            indent = ' ' * 4 * level
            doc.add_paragraph(f'{indent}{os.path.basename(root)}/', style=
                'Heading 3')
            subindent = ' ' * 4 * (level + 1)
            for file in files:
                if file.endswith('.py'):
                    file_path = os.path.join(root, file)
                    doc.add_paragraph(f'{subindent}{file}')
                    code_files.append(file_path)
    return code_files


def save_project_to_word(start_path, output_file):
    """
    Saves the project structure and code to a Word document, including only 'hermod/hermod'.

    Args:
        start_path (str): Path to the project directory.
        output_file (str): Output path for the Word document.
    """
    doc = Document()
    doc.add_heading('Project Code and Structure', 0)
    output_dir = os.path.dirname(output_file)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    code_files = write_directory_structure(doc, start_path)
    write_code_to_word(doc, code_files)
    doc.save(output_file)
    print(f'Word document saved: {output_file}')


home_directory = os.path.expanduser('~')
desktop_directory = os.path.join(home_directory, 'Desktop')
custom_output_path = os.path.join(desktop_directory, 'Sybertnetics',
    'sybertnetics model', 'hermod work')
if not os.path.exists(custom_output_path):
    os.makedirs(custom_output_path)
project_directory = os.path.abspath(os.path.dirname(__file__))
output_word_file = os.path.join(custom_output_path,
    'Project Structure And Code.docx')
print(f'Generating report for project: {project_directory}')
save_project_to_word(project_directory, output_word_file)
print(f'Report saved at: {output_word_file}')
print('Report generation complete!')
